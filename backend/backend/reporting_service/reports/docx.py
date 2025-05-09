
from django.http import HttpResponse
from docx import Document
from data_service.models import Schedule

def generate_schedule_docx():
    document = Document()
    document.add_heading('Расписание', level=1)
    
    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'День'
    hdr_cells[1].text = 'Время'
    hdr_cells[2].text = 'Дисциплина'
    hdr_cells[3].text = 'Преподаватель'
    hdr_cells[4].text = 'Аудитория'
    
    schedules = Schedule.objects.select_related(
        'teaching_load__discipline',
        'teaching_load__teacher',
        'classroom'
    ).all()
    
    for schedule in schedules:
        row_cells = table.add_row().cells
        row_cells[0].text = schedule.time_slot.get_day_of_week_display()
        row_cells[1].text = f"{schedule.time_slot.start_time}-{schedule.time_slot.end_time}"
        row_cells[2].text = schedule.teaching_load.discipline.name
        row_cells[3].text = schedule.teaching_load.teacher.short_name
        row_cells[4].text = schedule.classroom.number
    
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=schedule.docx'
    return response
